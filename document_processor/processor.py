"""
Document processor for extracting text, images, and structure from documents.
"""

import os
import tempfile
from typing import Dict, List, Any, Optional, Tuple, BinaryIO
import pytesseract
from PIL import Image
import pypdf
from docx import Document as DocxDocument
# Define our own simplified versions of partition functions
def partition_pdf(filename, **kwargs):
    """Simple implementation of partition_pdf."""
    pdf_reader = pypdf.PdfReader(filename)
    elements = []
    for i, page in enumerate(pdf_reader.pages):
        text = page.extract_text()
        if text:
            elements.append(type('Element', (), {'category': 'Text', 'text': text, '__str__': lambda self: self.text}))
    return elements

def partition_docx(filename, **kwargs):
    """Simple implementation of partition_docx."""
    doc = DocxDocument(filename)
    elements = []
    for para in doc.paragraphs:
        text = para.text
        if text:
            elements.append(type('Element', (), {'category': 'Text', 'text': text, '__str__': lambda self: self.text}))
    return elements
import openai
from pdf2image import convert_from_path

import sys
import os
# Add the parent directory to sys.path
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from config import (
    SUPPORTED_DOCUMENT_TYPES,
    MAX_IMAGE_SIZE,
    OCR_LANGUAGES,
    OPENAI_API_KEY,
    LLM_MODEL
)
from utils import (
    logger,
    get_file_extension,
    get_temp_file_path,
    create_document_metadata
)

class DocumentProcessor:
    """
    Process documents to extract text, images, and structure.
    """
    
    def __init__(self):
        """Initialize the document processor."""
        self.client = openai.OpenAI(api_key=OPENAI_API_KEY)
        logger.info("DocumentProcessor initialized")
    
    def process_document(self, file_path: str) -> Dict[str, Any]:
        """
        Process a document to extract text, images, and structure.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dict containing document metadata, text content, and image data
        """
        extension = get_file_extension(file_path)
        
        if extension not in SUPPORTED_DOCUMENT_TYPES:
            raise ValueError(f"Unsupported document type: {extension}")
        
        logger.info(f"Processing document: {file_path}")
        
        if extension in ['.pdf']:
            return self._process_pdf(file_path)
        elif extension in ['.docx', '.doc']:
            return self._process_docx(file_path)
        elif extension in ['.txt']:
            return self._process_txt(file_path)
        else:
            raise ValueError(f"Unsupported document type: {extension}")
    
    def _process_pdf(self, file_path: str) -> Dict[str, Any]:
        """Process a PDF document."""
        logger.info(f"Processing PDF: {file_path}")
        
        # Extract text and structure using unstructured
        elements = partition_pdf(
            filename=file_path,
            extract_images_in_pdf=True,
            infer_table_structure=True,
            chunking_strategy="by_title",
            max_characters=4000,
            new_after_n_chars=3800,
            combine_text_under_n_chars=2000,
        )
        
        # Extract images using pdf2image
        images_data = self._extract_images_from_pdf(file_path)
        
        # Get document metadata
        pdf_reader = pypdf.PdfReader(file_path)
        num_pages = len(pdf_reader.pages)
        title = os.path.basename(file_path)
        
        # Extract document structure using LLM
        sections = self._extract_document_structure(elements)
        
        # Create document metadata
        metadata = create_document_metadata(
            title=title,
            file_path=file_path,
            num_pages=num_pages,
            sections=sections,
            images=images_data
        )
        
        # Create full document content
        document_content = {
            "metadata": metadata,
            "elements": [self._element_to_dict(element) for element in elements],
            "images": images_data
        }
        
        logger.info(f"Processed PDF: {file_path}")
        return document_content
    
    def _process_docx(self, file_path: str) -> Dict[str, Any]:
        """Process a DOCX document."""
        logger.info(f"Processing DOCX: {file_path}")
        
        # Extract text and structure using unstructured
        elements = partition_docx(filename=file_path)
        
        # Extract images
        images_data = self._extract_images_from_docx(file_path)
        
        # Get document metadata
        doc = DocxDocument(file_path)
        num_pages = len(doc.paragraphs) // 40  # Approximate number of pages
        title = os.path.basename(file_path)
        
        # Extract document structure using LLM
        sections = self._extract_document_structure(elements)
        
        # Create document metadata
        metadata = create_document_metadata(
            title=title,
            file_path=file_path,
            num_pages=num_pages,
            sections=sections,
            images=images_data
        )
        
        # Create full document content
        document_content = {
            "metadata": metadata,
            "elements": [self._element_to_dict(element) for element in elements],
            "images": images_data
        }
        
        logger.info(f"Processed DOCX: {file_path}")
        return document_content
    
    def _process_txt(self, file_path: str) -> Dict[str, Any]:
        """Process a TXT document."""
        logger.info(f"Processing TXT: {file_path}")
        
        # Read text file
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        
        # Create elements
        elements = [{"type": "text", "text": text}]
        
        # Get document metadata
        num_pages = len(text) // 3000  # Approximate number of pages
        title = os.path.basename(file_path)
        
        # Extract document structure using LLM
        sections = self._extract_document_structure_from_text(text)
        
        # Create document metadata
        metadata = create_document_metadata(
            title=title,
            file_path=file_path,
            num_pages=max(1, num_pages),
            sections=sections,
            images=[]
        )
        
        # Create full document content
        document_content = {
            "metadata": metadata,
            "elements": elements,
            "images": []
        }
        
        logger.info(f"Processed TXT: {file_path}")
        return document_content
    
    def _extract_images_from_pdf(self, file_path: str) -> List[Dict[str, Any]]:
        """Extract images from a PDF document."""
        logger.info(f"Extracting images from PDF: {file_path}")
        
        images_data = []
        
        # Convert PDF pages to images
        images = convert_from_path(file_path)
        
        for i, image in enumerate(images):
            # Save image to temporary file
            temp_image_path = get_temp_file_path(f"pdf_page_{i}", ".jpg")
            image.save(temp_image_path)
            
            # Extract text from image using OCR
            ocr_text = pytesseract.image_to_string(image, lang="+".join(OCR_LANGUAGES))
            
            # Generate image description using LLM
            description = self._generate_image_description(image, ocr_text)
            
            # Add image data
            images_data.append({
                "page": i + 1,
                "path": temp_image_path,
                "ocr_text": ocr_text,
                "description": description
            })
        
        logger.info(f"Extracted {len(images_data)} images from PDF: {file_path}")
        return images_data
    
    def _extract_images_from_docx(self, file_path: str) -> List[Dict[str, Any]]:
        """Extract images from a DOCX document."""
        logger.info(f"Extracting images from DOCX: {file_path}")
        
        images_data = []
        doc = DocxDocument(file_path)
        
        # Create a temporary directory to extract images
        with tempfile.TemporaryDirectory() as temp_dir:
            # Process each paragraph to find images
            for i, para in enumerate(doc.paragraphs):
                for j, run in enumerate(para.runs):
                    # Check if run contains an image
                    if run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
                        # Extract image
                        for k, img in enumerate(run._element.findall('.//a:blip', namespaces={'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'})):
                            # Get image ID
                            img_id = img.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                            if img_id:
                                # Get image data
                                img_part = doc.part.related_parts[img_id]
                                img_bytes = img_part.blob
                                
                                # Save image to temporary file
                                temp_image_path = os.path.join(temp_dir, f"docx_image_{i}_{j}_{k}.jpg")
                                with open(temp_image_path, 'wb') as f:
                                    f.write(img_bytes)
                                
                                # Open image with PIL
                                image = Image.open(temp_image_path)
                                
                                # Extract text from image using OCR
                                ocr_text = pytesseract.image_to_string(image, lang="+".join(OCR_LANGUAGES))
                                
                                # Generate image description using LLM
                                description = self._generate_image_description(image, ocr_text)
                                
                                # Add image data
                                images_data.append({
                                    "paragraph": i,
                                    "path": temp_image_path,
                                    "ocr_text": ocr_text,
                                    "description": description
                                })
        
        logger.info(f"Extracted {len(images_data)} images from DOCX: {file_path}")
        return images_data
    
    def _generate_image_description(self, image: Image.Image, ocr_text: str) -> str:
        """Generate a description for an image using LLM."""
        # Resize image if necessary
        if image.width > MAX_IMAGE_SIZE[0] or image.height > MAX_IMAGE_SIZE[1]:
            image.thumbnail(MAX_IMAGE_SIZE)
        
        # Convert image to base64
        import io
        import base64
        buffered = io.BytesIO()
        image.save(buffered, format="JPEG")
        img_str = base64.b64encode(buffered.getvalue()).decode()
        
        # Generate description using OpenAI
        try:
            response = self.client.chat.completions.create(
                model=LLM_MODEL,
                messages=[
                    {"role": "system", "content": "You are an assistant that describes images for blind and low vision users. Provide detailed, concise descriptions that capture the key visual information."},
                    {"role": "user", "content": [
                        {"type": "text", "text": f"Describe this image in detail for a blind or low vision user. Include any text visible in the image. OCR text detected: {ocr_text}"},
                        {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{img_str}"}}
                    ]}
                ],
                max_tokens=300
            )
            description = response.choices[0].message.content
            return description
        except Exception as e:
            logger.error(f"Error generating image description: {e}")
            return f"Image with OCR text: {ocr_text}"
    
    def _extract_document_structure(self, elements: List[Any]) -> List[Dict[str, Any]]:
        """Extract document structure from elements using LLM."""
        # Convert elements to text
        elements_text = "\n\n".join([str(element) for element in elements])
        
        # Use LLM to extract structure
        return self._extract_document_structure_from_text(elements_text)
    
    def _extract_document_structure_from_text(self, text: str) -> List[Dict[str, Any]]:
        """Extract document structure from text using LLM."""
        # Truncate text if too long
        if len(text) > 10000:
            text = text[:10000] + "...[truncated]"
        
        try:
            # Use OpenAI to extract structure
            response = self.client.chat.completions.create(
                model=LLM_MODEL,
                messages=[
                    {"role": "system", "content": """You are an assistant that analyzes document structure. 
                     Extract the hierarchical structure of the document including sections, subsections, and their approximate locations.
                     Format your response as a JSON array of sections, where each section has:
                     - title: The section title
                     - level: The heading level (1 for main sections, 2 for subsections, etc.)
                     - start_position: Approximate position in the document (e.g., "beginning", "middle", "end")
                     - content_summary: A brief summary of the section content"""},
                    {"role": "user", "content": f"Analyze the structure of this document:\n\n{text}"}
                ],
                response_format={"type": "json_object"},
                max_tokens=1000
            )
            
            structure_text = response.choices[0].message.content
            import json
            structure = json.loads(structure_text)
            
            # Ensure we have a "sections" key
            if "sections" in structure:
                return structure["sections"]
            else:
                return structure.get("sections", [])
                
        except Exception as e:
            logger.error(f"Error extracting document structure: {e}")
            # Return a basic structure
            return [{"title": "Document", "level": 1, "start_position": "beginning", "content_summary": "Document content"}]
    
    def _element_to_dict(self, element: Any) -> Dict[str, Any]:
        """Convert an unstructured element to a dictionary."""
        element_dict = {
            "type": element.category if hasattr(element, "category") else type(element).__name__,
            "text": str(element)
        }
        
        # Add additional properties if available
        if hasattr(element, "metadata"):
            element_dict["metadata"] = element.metadata
        
        return element_dict
